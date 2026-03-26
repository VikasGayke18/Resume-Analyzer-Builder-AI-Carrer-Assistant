from flask import Flask, render_template, request, redirect, url_for, session, send_file, flash
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from datetime import datetime
from io import BytesIO
import os
import json
import re
from interview_questions import INTERVIEW_QUESTIONS
from expected_answers import EXPECTED_ANSWERS
from history_manager import save_interview
import random
from flask import jsonify
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.lib import colors
from io import BytesIO
from reportlab.platypus import Table, TableStyle, Image
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
    HRFlowable
)
import pdfkit
import base64
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

app = Flask(__name__)
app.secret_key = "your_secret_key_here"
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///career_ai.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = "uploads"

db = SQLAlchemy(app)

# Database Models

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Activity(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    activity_type = db.Column(db.String(100), nullable=False)   # analyze, builder, chatbot, interview
    title = db.Column(db.String(255), nullable=False)
    details = db.Column(db.Text)
    timestamp = db.Column(db.DateTime, default=datetime.utcnow)

class InterviewSession(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)
    role = db.Column(db.String(100), nullable=False)
    level = db.Column(db.String(50), nullable=False)
    score = db.Column(db.Integer, default=0)
    total_questions = db.Column(db.Integer, default=0)
    correct_answers = db.Column(db.Integer, default=0)
    transcript = db.Column(db.Text)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

# Utility

if not os.path.exists(app.config["UPLOAD_FOLDER"]):
    os.makedirs(app.config["UPLOAD_FOLDER"])

def login_required():
    return "user_id" in session

def save_activity(user_id, activity_type, title, details=""):
    activity = Activity(
        user_id=user_id,
        activity_type=activity_type,
        title=title,
        details=details
    )
    db.session.add(activity)
    db.session.commit()


def extract_text_from_pdf(filepath):
    # basic safe placeholder text extraction
    # you can replace with PyPDF2/pdfplumber later
    try:
        import PyPDF2
        text = ""
        with open(filepath, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + " "
        return text.lower()
    except Exception:
        return ""

def chatbot_reply(message):

    msg = message.lower()

    # GET ANALYZER DATA
    results = session.get("last_analysis")

    #  PERSONALIZED RESPONSES

    if results:

        if "score" in msg:
            return f"Your current ATS score is {results.get('score',0)}%. Try improving keywords and sections."

        elif "missing skills" in msg or "missing keywords" in msg:
            missing = results.get("missing", [])
            if missing:
                return "You are missing these important skills:\n- " + "\n- ".join(missing)
            else:
                return "Great! You are not missing any major keywords."

        elif "keywords" in msg:
            found = results.get("found", [])
            return "These keywords were found:\n- " + "\n- ".join(found)

        elif "sections" in msg:
            missing_sec = results.get("missing_sections", [])
            if missing_sec:
                return "You should add these sections:\n- " + "\n- ".join(missing_sec)
            else:
                return "Awesome! All important sections are present."

        elif "action words" in msg:
            actions = results.get("action_words_used", [])
            if actions:
                return "You used these action words:\n- " + "\n- ".join(actions)
            else:
                return "Try using words like Developed, Built, Designed."

        elif "improve" in msg:
            suggestions = results.get("suggestions", [])
            if suggestions:
                return "Here’s how you can improve:\n- " + "\n- ".join(suggestions)
            else:
                return "Your resume looks strong already!"

    
    # 🤖 GENERAL CHATBOT

    if "hello" in msg or "hi" in msg or "hey" in msg:
        return "Hey! 👋 Ask me about your resume analysis or career tips."

    elif "resume" in msg:
        return "A good resume should have strong keywords, proper sections, and clean formatting."

    else:
        return f"""I understand your question: "{message}"

General tips:
- Add relevant keywords
- Use bullet points
- Add measurable achievements
- Keep it clean and structured"""

def generate_ai_response(message):

    msg = message.lower()

    # KEYWORD BASED INTENTS

    if "score" in msg:
        return "Your ATS score depends on keywords, formatting, and sections. Try adding relevant skills and proper structure."

    elif "improve" in msg or "better" in msg:
        return """To improve your resume:
- Add role-specific keywords
- Use bullet points
- Add measurable achievements (e.g., increased performance by 20%)
- Keep format clean and structured"""

    elif "missing skills" in msg or "skills" in msg:
        return "Check the 'Missing Keywords' section in analyzer and try adding those skills in your resume."

    elif "project" in msg:
        return "Add strong projects with clear description, technologies used, and results."

    elif "experience" in msg:
        return "Mention your work clearly using action verbs like Developed, Built, Designed."

    elif "hello" in msg or "hi" in msg or "hey" in msg:
        return "Hey! 👋 Ask me anything about resumes, ATS score, or career tips."

    elif "ats" in msg:
        return "ATS checks keywords, formatting, sections, and readability. Optimize all for better score."

    else:
        #  FALLBACK (SMART DEFAULT)
        return f"""I understand your question: "{message}"

Here are some general tips:
- Improve resume structure
- Add relevant keywords
- Use bullet points
- Include measurable achievements"""

INTERVIEW_QUESTIONS = {

"Software Developer":{

"basic": [
{"q":"What is Python?","keywords":["programming","language"],"answer":"Python is a high-level programming language that is known for its simplicity and readability. It is widely used as a programming language in various fields such as web development, data science, and artificial intelligence. This language allows developers to write clean and efficient code. Because of its easy syntax, it is one of the most popular programming languages for beginners."},
{"q":"What is a variable?","keywords":["store","data"],"answer":"A variable is used in programming to store data values that can be used later in a program. It acts like a container where you can store different types of data such as numbers, text, or lists. The value of a variable can change during the execution of a program. This ability to store and modify data makes variables very important in coding."},
{"q":"What is a function?","keywords":["reusable","code"],"answer":"A function is a reusable block of code that is designed to perform a specific task. Instead of writing the same code again and again, you can call a function whenever needed. This makes the code more organized and easier to maintain. Functions help improve efficiency by allowing developers to reuse code in different parts of a program."},
{"q":"What is Git?","keywords":["version","control"],"answer":"Git is a version control system that helps developers track changes in their code over time. It allows multiple people to work on the same project without conflicts. With version control, you can easily go back to previous versions of your code if needed. Git is widely used in software development for collaboration and managing code history."},
{"q":"What is debugging?","keywords":["error","fix"],"answer":"Debugging is the process of identifying and fixing an error in a program. When a program does not work as expected, developers use debugging techniques to find the issue. It helps in understanding why the error occurred and how to fix it properly. Debugging is an essential skill for programmers to ensure their code runs correctly."},
{"q":"What is HTML?","keywords":["web","structure"],"answer":"HTML is a markup language used to create the structure of web pages. It defines elements like headings, paragraphs, images, and links on the web. This structure helps browsers display content in a proper format. HTML is the foundation of every website and works together with other technologies to build web applications."},
{"q":"What is CSS?","keywords":["style","design"],"answer":"CSS is used to add style and design to web pages created with HTML. It controls how elements look, including colors, fonts, spacing, and layout. With CSS, developers can make web pages more attractive and user-friendly. It plays a key role in improving the visual appearance and overall design of a website."},
{"q":"What is JavaScript?","keywords":["interactive","web"],"answer":"JavaScript is a programming language used to make web pages interactive and dynamic. It allows developers to add features like buttons, animations, and real-time updates on the web. This interactivity improves user experience significantly. JavaScript works along with HTML and CSS to create modern and responsive websites."},
{"q":"What is a loop?","keywords":["repeat","iteration"],"answer":"A loop is used in programming to repeat a block of code multiple times until a condition is met. It helps automate repetitive tasks and reduces the need to write the same code again. Each cycle of a loop is called an iteration. Loops are very useful when working with large amounts of data or repeated operations."},
{"q":"What is an array?","keywords":["collection","data"],"answer":"An array is a collection of multiple data elements stored in a single variable. It allows you to manage and access related data efficiently. Each element in the array can be accessed using an index. Arrays are widely used in programming to handle and organize data in a structured way."}
],

"intermediate": [
{"q":"What is OOP?","keywords":["class","object"],"answer":"Object Oriented Programming (OOP) is a programming approach based on the concept of class and object. A class defines the structure and behavior, while an object is an instance of that class. It helps organize code in a more modular and reusable way. OOP includes principles like inheritance, encapsulation, and polymorphism to improve code efficiency and maintainability."},
{"q":"What is API?","keywords":["communication","interface"],"answer":"An API (Application Programming Interface) is a communication interface that allows different software applications to interact with each other. It acts as a bridge enabling smooth communication between systems by sending and receiving data. APIs define rules and protocols that developers follow to connect applications. In simple terms, it is an interface that enables communication between different software components."},
{"q":"What is REST API?","keywords":["http","request"],"answer":"A REST API is a type of API that follows REST principles and uses HTTP methods for communication. It works by sending an HTTP request from the client to the server and receiving a response. Common methods include GET, POST, PUT, and DELETE. REST APIs are widely used because they are simple, scalable, and efficient for web services."},
{"q":"What is SQL?","keywords":["database","query"],"answer":"SQL (Structured Query Language) is used to manage and manipulate data in a database. It allows users to write a query to retrieve, insert, update, or delete data. SQL is essential for working with relational database systems. It provides a standard way to interact with structured data efficiently."},
{"q":"What is MVC?","keywords":["model","view","controller"],"answer":"MVC is a software design pattern that divides an application into three parts: model, view, and controller. The model handles the data and business logic, the view manages the user interface, and the controller processes user input. This separation makes the code more organized and easier to manage. It also improves scalability and maintainability of applications."},
{"q":"What is JSON?","keywords":["data","format"],"answer":"JSON (JavaScript Object Notation) is a lightweight data format used for storing and exchanging data. It is easy to read and write for humans and simple for machines to parse. JSON represents data in key-value pairs, making it structured and organized. It is commonly used in APIs for data transfer between client and server."},
{"q":"What is error handling?","keywords":["exception","try"],"answer":"Error handling is the process of managing runtime issues in a program using techniques like exception handling. It uses constructs such as try and catch to detect and handle errors gracefully. This prevents the program from crashing unexpectedly. Proper error handling helps developers debug and maintain reliable applications."},
{"q":"What is a class?","keywords":["blueprint","object"],"answer":"A class is a blueprint used to create objects in programming. It defines the properties and behaviors that an object will have. By using a class, developers can create multiple objects with similar characteristics. This helps in organizing code and improving reusability."},
{"q":"What is inheritance?","keywords":["parent","child"],"answer":"Inheritance is a feature in OOP where a child class can inherit properties and methods from a parent class. This helps in reusing existing code and reducing duplication. The child class can also add new features or modify existing ones. It improves the structure and hierarchy of code."},
{"q":"What is encapsulation?","keywords":["data","hiding"],"answer":"Encapsulation is a concept in OOP that focuses on data hiding and restricting direct access to data. It is achieved by using methods to control how data is accessed or modified. This protects the integrity of the data and prevents unintended changes. Encapsulation improves security and maintainability of the code."}
],

"advanced": [
{"q":"What is system design?","keywords":["architecture","scalable"],"answer":"System design is the process of defining the architecture of a system to meet specific requirements. It focuses on building scalable and efficient systems that can handle real-world demands. A good system design ensures reliability, performance, and maintainability. It involves planning components, data flow, and interactions between different parts of the system."},
{"q":"What is microservices architecture?","keywords":["services","distributed"],"answer":"Microservices architecture is a design approach where an application is divided into small independent services. Each of these services performs a specific function and can be developed and deployed separately. It is a distributed system where services communicate with each other through APIs. This approach improves flexibility, scalability, and fault isolation."},
{"q":"What is load balancing?","keywords":["traffic","server"],"answer":"Load balancing is a technique used to distribute incoming traffic across multiple server systems. It helps ensure that no single server becomes overloaded, improving performance and reliability. By balancing traffic efficiently, it enhances user experience and reduces downtime. Load balancing is essential in large-scale applications to maintain system stability."},
{"q":"What is caching?","keywords":["memory","performance"],"answer":"Caching is a technique used to store frequently accessed data in memory for faster retrieval. It reduces the need to repeatedly fetch data from slower storage systems. By using caching, applications can significantly improve performance and response time. It is widely used in web applications and databases to optimize speed."},
{"q":"What is multithreading?","keywords":["parallel","threads"],"answer":"Multithreading is a programming concept that allows multiple threads to run in parallel within a single program. Each thread performs a separate task, improving efficiency and execution speed. It helps in better utilization of CPU resources. Multithreading is commonly used in applications that require concurrent processing."},
{"q":"What is Docker?","keywords":["container","deployment"],"answer":"Docker is a platform used to create, run, and manage applications inside a container. A container packages the application along with its dependencies, ensuring consistency across environments. This makes deployment faster and more reliable. Docker simplifies the development and deployment process in modern software systems."},
{"q":"What is Kubernetes?","keywords":["container","orchestration"],"answer":"Kubernetes is a powerful platform used for container orchestration. It manages multiple containerized applications across a cluster of machines. Kubernetes automates tasks like deployment, scaling, and monitoring of containers. It is widely used for handling large-scale applications efficiently."},
{"q":"What is CI/CD?","keywords":["automation","deployment"],"answer":"CI/CD stands for Continuous Integration and Continuous Deployment, which focuses on automation in the software development lifecycle. It allows developers to automatically test and deploy code changes. This reduces manual effort and speeds up the deployment process. CI/CD helps maintain code quality and ensures faster delivery of applications."},
{"q":"What is scalability?","keywords":["increase","performance"],"answer":"Scalability refers to the ability of a system to handle an increase in workload without affecting performance. A scalable system can grow by adding more resources like servers or processing power. It ensures that applications continue to perform well under heavy load. Scalability is a key factor in designing modern systems."},
{"q":"What is distributed system?","keywords":["multiple","nodes"],"answer":"A distributed system is a system where multiple computers, also known as nodes, work together to achieve a common goal. These nodes communicate and coordinate with each other over a network. It helps improve performance, reliability, and scalability. Distributed systems are widely used in large-scale applications like cloud computing."}
]
},

"Data Analyst": {

"basic":[
{"q":"What is data analysis?","keywords":["data","insight"],"answer":"Data analysis is the process of examining and interpreting data to extract meaningful insight. It involves collecting, organizing, and analyzing data to support decision-making. By using various tools and techniques, analysts can identify patterns and trends. This helps businesses and individuals make informed decisions based on data."},
{"q":"What is Excel?","keywords":["spreadsheet","data"],"answer":"Excel is a powerful spreadsheet tool used for organizing and analyzing data. It allows users to perform calculations, create charts, and manage large amounts of data efficiently. With features like formulas and pivot tables, it simplifies complex tasks. Excel is widely used in business, finance, and data analysis."},
{"q":"What is dataset?","keywords":["collection","data"],"answer":"A dataset is a collection of related data that is organized for analysis. It can include numbers, text, or other types of data arranged in a structured format. Datasets are commonly used in research, machine learning, and business analytics. Proper organization of data in a dataset helps in extracting useful information."},
{"q":"What is data cleaning?","keywords":["remove","errors"],"answer":"Data cleaning is the process used to remove errors and inconsistencies from data. It ensures that the data is accurate, complete, and reliable for analysis. This may include fixing missing values, correcting incorrect entries, and standardizing formats. Clean data is essential for generating accurate results and insights."},
{"q":"What is CSV file?","keywords":["comma","data"],"answer":"A CSV (Comma-Separated Values) file is a simple format used to store tabular data. In this format, each value is separated by a comma, making it easy to read and write. CSV files are widely used for data exchange between different systems. They are lightweight and supported by many tools and programming languages."},
{"q":"What is row and column?","keywords":["table","data"],"answer":"In a table, data is organized into rows and columns for better structure. A row represents a single record, while a column represents a specific attribute of that data. This format helps in easy understanding and analysis of data. Rows and columns are fundamental concepts in databases and spreadsheets."},
{"q":"What is data type?","keywords":["integer","string"],"answer":"A data type defines the kind of data that can be stored and processed in a program. Common types include integer for numbers and string for text. Data types help the system understand how to handle different kinds of data. Using the correct data type ensures accuracy and efficiency in data processing."},
{"q":"What is filtering?","keywords":["select","data"],"answer":"Filtering is a technique used to select specific data from a dataset based on certain conditions. It helps in focusing only on the relevant data needed for analysis. By applying filters, users can quickly find important information. Filtering is commonly used in spreadsheets and databases."},
{"q":"What is sorting?","keywords":["order","data"],"answer":"Sorting is the process of arranging data in a specific order, such as ascending or descending. It helps in organizing data for easier understanding and analysis. Sorting can be applied to numbers, text, or dates. This makes it easier to identify patterns and compare values."},
{"q":"What is dashboard?","keywords":["visual","data"],"answer":"A dashboard is a tool used to display data in a visual format using charts, graphs, and tables. It helps users quickly understand complex data through visual representation. Dashboards are commonly used in business intelligence and analytics. They provide real-time insights and support better decision-making."}
],

"intermediate":[
{"q":"What is Pandas?","keywords":["python","dataframe"],"answer":"Pandas is a powerful python library used for data manipulation and analysis. It provides data structures like dataframe that make it easy to work with structured data. With Pandas, users can clean, filter, and transform data efficiently. It is widely used in data science and analytics for handling large datasets."},
{"q":"What is NumPy?","keywords":["array","python"],"answer":"NumPy is a fundamental python library used for numerical computations. It provides support for large multi-dimensional array objects and matrices. NumPy allows fast and efficient operations on array data. It is commonly used as the base for many other scientific and data analysis libraries."},
{"q":"What is SQL?","keywords":["database","query"],"answer":"SQL (Structured Query Language) is used to manage and interact with a database. It allows users to write a query to retrieve, insert, update, or delete data. SQL is essential for handling structured data stored in relational databases. It provides a standard way to work with database systems."},
{"q":"What is data visualization?","keywords":["charts","graphs"],"answer":"Data visualization is the process of representing data using charts and graphs. It helps in understanding complex data by presenting it in a visual format. Visual tools like bar charts, line graphs, and pie charts make analysis easier. It is widely used in business and analytics to communicate insights clearly."},
{"q":"What is regression?","keywords":["prediction","model"],"answer":"Regression is a statistical method used for prediction of continuous values. It builds a model based on the relationship between variables in data. This model helps in forecasting trends and outcomes. Regression is commonly used in machine learning and data analysis."},
{"q":"What is correlation?","keywords":["relationship","data"],"answer":"Correlation is a statistical measure that shows the relationship between two variables in data. It indicates how strongly the variables are connected. A positive or negative value represents the direction of the relationship. Correlation helps in understanding patterns and dependencies in datasets."},
{"q":"What is ETL?","keywords":["extract","transform","load"],"answer":"ETL stands for extract, transform, and load, which is a process used in data integration. First, data is extracted from different sources. Then it is transformed into a suitable format for analysis. Finally, the data is loaded into a system like a data warehouse for further use."},
{"q":"What is aggregation?","keywords":["sum","average"],"answer":"Aggregation is the process of summarizing data using functions like sum and average. It helps in combining multiple data values into a single result. Aggregation is useful for analyzing trends and patterns in large datasets. It is commonly used in databases and data analysis tools."},
{"q":"What is group by?","keywords":["categorize","data"],"answer":"Group by is a technique used to categorize data based on one or more columns. It groups similar data together for easier analysis. After grouping, operations like sum or count can be applied to each group. It is widely used in SQL and data analysis libraries like Pandas."},
{"q":"What is missing data?","keywords":["null","values"],"answer":"Missing data refers to null or undefined values present in a dataset. These values can affect the accuracy of analysis if not handled properly. Techniques like removal or imputation are used to deal with missing values. Handling missing data is an important step in data preprocessing."}
],

"advanced":[
{"q":"What is big data?","keywords":["large","data"],"answer":"Big data refers to extremely large volumes of data that cannot be processed using traditional methods. This large data comes from various sources like social media, sensors, and transactions. It requires advanced tools and technologies for storage and analysis. Big data helps organizations gain insights and make better decisions."},
{"q":"What is data pipeline?","keywords":["flow","process"],"answer":"A data pipeline is a system that manages the flow of data from one place to another. It automates the process of collecting, transforming, and delivering data. This process ensures that data is available for analysis in the right format. Data pipelines are essential for handling continuous data processing efficiently."},
{"q":"What is A/B testing?","keywords":["experiment","compare"],"answer":"A/B testing is an experiment where two versions of something are compared to determine which one performs better. It is commonly used in marketing, websites, and product design. By comparing user responses, businesses can make data-driven decisions. This method helps improve performance and user experience."},
{"q":"What is feature engineering?","keywords":["transform","data"],"answer":"Feature engineering is the process of using techniques to transform data into a better format for machine learning models. It involves selecting, modifying, or creating new features from raw data. This transformation improves the accuracy and performance of models. It is a crucial step in building effective predictive systems."},
{"q":"What is data warehouse?","keywords":["storage","data"],"answer":"A data warehouse is a system used for storage of large amounts of structured data. It collects data from multiple sources and organizes it for analysis. Data warehouses are optimized for querying and reporting. They help businesses make strategic decisions using historical data."},
{"q":"What is data lake?","keywords":["raw","data"],"answer":"A data lake is a storage system that holds raw data in its original format. This raw data can be structured, semi-structured, or unstructured. It provides flexibility for storing large volumes of information. Data lakes are commonly used in big data and machine learning applications."},
{"q":"What is predictive analysis?","keywords":["forecast","data"],"answer":"Predictive analysis is a technique that uses historical data to forecast future outcomes. It applies statistical models and machine learning algorithms to identify patterns. This helps organizations make informed predictions. Predictive analysis is widely used in finance, marketing, and healthcare."},
{"q":"What is clustering?","keywords":["group","data"],"answer":"Clustering is a technique used to group similar data points together based on their characteristics. It is commonly used in machine learning and data analysis. By grouping data, it becomes easier to identify patterns and trends. Clustering helps in segmentation and understanding large datasets."},
{"q":"What is anomaly detection?","keywords":["outlier","data"],"answer":"Anomaly detection is the process of identifying unusual patterns or outlier data points. These outliers may indicate errors, fraud, or rare events. It is widely used in cybersecurity, finance, and monitoring systems. Detecting anomalies helps improve system reliability and security."},
{"q":"What is KPI?","keywords":["performance","metric"],"answer":"KPI (Key Performance Indicator) is a performance metric used to evaluate the success of an organization or activity. It helps track progress towards specific goals. KPIs can be financial, operational, or customer-related. They are essential for measuring and improving performance over time."}
]

},

"AI Engineer": {

"basic":[
{"q":"What is AI?","keywords":["machine","intelligence"],"answer":"AI, or Artificial Intelligence, is the simulation of human intelligence in a machine. It enables machines to think, learn, and make decisions like humans. AI systems use data and algorithms to perform tasks such as problem-solving and pattern recognition. This intelligence in machines is widely used in various industries like healthcare, finance, and automation."},
{"q":"What is machine learning?","keywords":["model","data"],"answer":"Machine learning is a subset of AI that allows systems to learn from data without being explicitly programmed. It uses algorithms to build a model that can make predictions or decisions. By training on data, the model improves its performance over time. Machine learning is widely used in recommendation systems, image recognition, and more."},
{"q":"What is dataset?","keywords":["training","data"],"answer":"A dataset is a collection of data used for training machine learning models. It contains input and sometimes output values that help the model learn patterns. The quality and size of training data directly affect model performance. Datasets are essential for building accurate and reliable machine learning systems."},
{"q":"What is training data?","keywords":["learn","model"],"answer":"Training data is the data used to help a model learn patterns and relationships. It provides examples that the model uses to understand how inputs relate to outputs. The better the quality of training data, the better the model can learn. Training is a key step in building any machine learning system."},
{"q":"What is testing data?","keywords":["evaluate","model"],"answer":"Testing data is used to evaluate how well a trained model performs on unseen data. It helps measure the accuracy and effectiveness of the model. This data is separate from the training data to ensure fair evaluation. Testing is important to check if the model can generalize well."},
{"q":"What is feature?","keywords":["input","data"],"answer":"A feature is an individual input variable used in a machine learning model. It represents a specific aspect of the data that helps in making predictions. Features are selected and processed carefully to improve model performance. Good quality input data leads to better results in machine learning."},
{"q":"What is label?","keywords":["output","data"],"answer":"A label is the output data or target variable in supervised learning. It represents the correct answer that the model is trying to predict. During training, the model learns the relationship between input data and labels. Labels are essential for guiding the learning process."},
{"q":"What is model?","keywords":["algorithm","prediction"],"answer":"A model is a trained algorithm that is used for making prediction based on input data. It learns patterns from training data and applies them to new data. Models can be simple or complex depending on the problem. They are the core component of machine learning systems."},
{"q":"What is algorithm?","keywords":["steps","solution"],"answer":"An algorithm is a set of defined steps used to solve a problem or perform a task. It provides a clear solution by following a sequence of operations. In machine learning, algorithms are used to train models and make predictions. Efficient algorithms improve performance and accuracy."},
{"q":"What is Python in AI?","keywords":["programming","ai"],"answer":"Python is a popular programming language widely used in AI development. It provides many libraries and tools for building and training models. Python is easy to learn and has strong community support. Because of its simplicity and power, it is the preferred programming language for AI and machine learning projects."}
],

"intermediate":[
{"q":"What is neural network?","keywords":["neurons","layers"],"answer":"A neural network is a machine learning model inspired by the human brain. It is made up of interconnected neurons organized into different layers. These layers process input data and pass information forward to generate outputs. Neural networks are widely used in tasks like image recognition and speech processing."},
{"q":"What is deep learning?","keywords":["neural","layers"],"answer":"Deep learning is a subset of machine learning that uses neural networks with multiple layers. These layers allow the model to learn complex patterns from large amounts of data. Deep learning is widely used in applications like image recognition and natural language processing. It provides powerful solutions for complex problems."},
{"q":"What is overfitting?","keywords":["training","error"],"answer":"Overfitting occurs when a model learns the training data too well, including noise and unnecessary details. This results in very low error on training data but poor performance on new data. It means the model is not generalizing well. Techniques like regularization and cross-validation help reduce overfitting."},
{"q":"What is underfitting?","keywords":["model","poor"],"answer":"Underfitting happens when a model is too simple to capture the underlying patterns in data. As a result, it performs poorly on both training and test data. This indicates that the model has not learned enough from the data. Improving model complexity can help reduce underfitting."},
{"q":"What is NLP?","keywords":["text","language"],"answer":"NLP, or Natural Language Processing, is a field of AI that focuses on understanding and processing human language. It enables machines to analyze text and extract meaningful information. NLP is used in applications like chatbots, translation, and sentiment analysis. It helps bridge the gap between human language and computers."},
{"q":"What is TensorFlow?","keywords":["deep","learning"],"answer":"TensorFlow is an open-source library used for deep learning and machine learning tasks. It provides tools to build, train, and deploy models efficiently. TensorFlow supports large-scale computation and works well with neural networks. It is widely used in both research and production environments."},
{"q":"What is PyTorch?","keywords":["framework","ai"],"answer":"PyTorch is a popular framework used for building AI and deep learning models. It provides flexibility and ease of use for developers and researchers. PyTorch supports dynamic computation graphs, making it easier to debug and experiment. It is widely used in modern AI applications."},
{"q":"What is classification?","keywords":["categories","predict"],"answer":"Classification is a type of machine learning task where the goal is to predict categories of data. The model learns from labeled data and assigns new inputs to predefined classes. It is commonly used in spam detection, image classification, and medical diagnosis. Classification helps in making decisions based on categories."},
{"q":"What is regression?","keywords":["continuous","predict"],"answer":"Regression is a machine learning technique used to predict continuous values. It models the relationship between input variables and output data. Regression is widely used in forecasting and trend analysis. It helps in predicting values like prices, temperature, and demand."},
{"q":"What is feature selection?","keywords":["important","data"],"answer":"Feature selection is the process of selecting the most important features from data for model training. It helps reduce complexity and improve model performance. By removing irrelevant features, the model becomes faster and more accurate. Feature selection is an important step in data preprocessing."}
],

"advanced":[
{"q":"What is transformer model?","keywords":["attention","nlp"],"answer":"A transformer model is a type of deep learning model widely used in NLP tasks. It is based on the attention mechanism, which helps the model focus on important parts of the input data. Transformers process data in parallel, making them faster than traditional models. They are commonly used in applications like translation, chatbots, and text generation."},
{"q":"What is generative AI?","keywords":["generate","content"],"answer":"Generative AI is a type of artificial intelligence that can generate new content such as text, images, or audio. It learns patterns from existing data and creates similar but original outputs. This technology is used in chatbots, image generation, and creative applications. Generative AI is rapidly growing and transforming many industries."},
{"q":"What is reinforcement learning?","keywords":["reward","agent"],"answer":"Reinforcement learning is a machine learning approach where an agent learns by interacting with an environment. The agent receives a reward for correct actions and penalties for wrong ones. Over time, it learns the best strategy to maximize rewards. This technique is widely used in robotics, gaming, and decision-making systems."},
{"q":"What is model optimization?","keywords":["performance","tuning"],"answer":"Model optimization is the process of improving the performance of a machine learning model. It involves tuning different aspects of the model to achieve better accuracy and efficiency. Optimization techniques help reduce errors and improve predictions. It is an important step in building high-quality models."},
{"q":"What is hyperparameter tuning?","keywords":["parameters","model"],"answer":"Hyperparameter tuning is the process of selecting the best parameters for a model to improve its performance. These parameters are set before training and control how the model learns. Proper tuning can significantly improve accuracy and efficiency. It is commonly done using techniques like grid search or random search."},
{"q":"What is bias in AI?","keywords":["error","model"],"answer":"Bias in AI refers to systematic error in a model that leads to incorrect predictions. It occurs when the model makes assumptions that oversimplify the problem. High bias can cause underfitting and poor performance. Reducing bias is important for building fair and accurate AI systems."},
{"q":"What is variance?","keywords":["overfit","model"],"answer":"Variance refers to how much a model's predictions change with different training data. A model with high variance tends to overfit the data and perform poorly on new data. It captures noise instead of actual patterns. Balancing variance is important for creating a generalizable model."},
{"q":"What is cross validation?","keywords":["split","data"],"answer":"Cross validation is a technique used to evaluate the performance of a model by splitting data into multiple parts. The model is trained and tested on different subsets of the data. This helps ensure that the model performs well on unseen data. It is widely used to improve reliability and avoid overfitting."},
{"q":"What is GAN?","keywords":["generate","network"],"answer":"GAN, or Generative Adversarial Network, is a type of neural network used to generate new data. It consists of two networks: a generator and a discriminator. The generator tries to generate realistic data, while the discriminator evaluates it. This process improves the quality of generated data over time."},
{"q":"What is embedding?","keywords":["vector","representation"],"answer":"Embedding is a technique used to convert data into a vector representation. It helps represent complex data like text or images in numerical form. These vector representations capture relationships and patterns in data. Embeddings are widely used in NLP, recommendation systems, and deep learning models."}
]
},

"Cyber Security": {

"basic":[
{"q":"What is cyber security?","keywords":["protect","systems"],"answer":"Cyber security is the practice used to protect systems, networks, and data from cyber threats and attacks. It involves using tools and techniques to secure digital information. Cyber security helps prevent unauthorized access and data breaches. It is essential for maintaining the safety and privacy of systems in the digital world."},
{"q":"What is a virus?","keywords":["malware","harm"],"answer":"A virus is a type of malware that is designed to harm computers and spread from one system to another. It attaches itself to files or programs and activates when executed. Viruses can corrupt data, slow down systems, or cause system failure. Preventing malware infections is important for system security."},
{"q":"What is firewall?","keywords":["network","security"],"answer":"A firewall is a network security system that monitors and controls incoming and outgoing traffic. It acts as a barrier between trusted and untrusted networks. Firewalls use predefined rules to allow or block data packets. They are essential for protecting systems from unauthorized access."},
{"q":"What is password security?","keywords":["strong","authentication"],"answer":"Password security involves creating and maintaining strong passwords to protect user accounts. A strong password includes a mix of letters, numbers, and special characters. It is an important part of authentication to verify user identity. Proper password practices help prevent unauthorized access."},
{"q":"What is phishing?","keywords":["fake","email"],"answer":"Phishing is a type of cyber attack where attackers use fake email messages to trick users into sharing sensitive information. These emails often appear to be from trusted sources. Users may unknowingly provide passwords or financial details. Awareness of fake email scams is important for protection."},
{"q":"What is antivirus?","keywords":["protect","software"],"answer":"Antivirus is a type of software designed to protect systems from malware and other threats. It scans files and programs to detect harmful activities. Antivirus software can remove or quarantine infected files. Regular updates ensure better protection against new threats."},
{"q":"What is encryption?","keywords":["secure","data"],"answer":"Encryption is the process of converting data into a secure format to prevent unauthorized access. Only authorized users with the correct key can read the encrypted data. It is widely used to protect sensitive information during transmission. Encryption ensures data privacy and security."},
{"q":"What is hacker?","keywords":["access","system"],"answer":"A hacker is a person who attempts to gain access to a system or network, often without permission. Some hackers use their skills for malicious purposes, while others work ethically to improve security. Unauthorized access can lead to data breaches and damage. Understanding hacking helps in improving cyber security measures."},
{"q":"What is malware?","keywords":["malicious","software"],"answer":"Malware is malicious software designed to damage, disrupt, or gain unauthorized access to systems. It includes viruses, worms, ransomware, and spyware. Malware can steal data, slow down systems, or cause failures. Protecting systems from malicious threats is essential for safety."},
{"q":"What is authentication?","keywords":["verify","identity"],"answer":"Authentication is the process used to verify the identity of a user before granting access to a system. It ensures that only authorized users can access sensitive data. Common methods include passwords, biometrics, and OTPs. Authentication is a key part of system security."}
],

"intermediate":[
{"q":"What is penetration testing?","keywords":["test","security"],"answer":"Penetration testing is a method used to test the security of a system by simulating real-world attacks. It helps identify vulnerabilities before attackers can exploit them. Security professionals perform this test using various tools and techniques. The main goal is to improve system security and protect sensitive data."},
{"q":"What is SQL injection?","keywords":["attack","database"],"answer":"SQL injection is a type of cyber attack that targets a database by inserting malicious SQL queries. These queries can manipulate or access sensitive data without authorization. It is one of the most common web vulnerabilities. Proper input validation and security measures can prevent such attacks."},
{"q":"What is XSS?","keywords":["script","attack"],"answer":"Cross-Site Scripting (XSS) is a type of attack where malicious script code is injected into web pages. This script can run in the browser of other users and steal data or perform unwanted actions. XSS attacks exploit vulnerabilities in web applications. Proper validation and sanitization of input can help prevent such attacks."},
{"q":"What is VPN?","keywords":["secure","network"],"answer":"A VPN, or Virtual Private Network, is used to create a secure connection over the internet. It encrypts data and hides the user's identity while accessing a network. VPNs are commonly used to protect privacy and access restricted content. They provide an extra layer of security for online activities."},
{"q":"What is IDS?","keywords":["detect","intrusion"],"answer":"An Intrusion Detection System (IDS) is a security tool used to detect suspicious activities or intrusion attempts in a network. It monitors traffic and identifies potential threats. IDS alerts administrators when unusual behavior is detected. It plays an important role in maintaining network security."},
{"q":"What is IPS?","keywords":["prevent","attack"],"answer":"An Intrusion Prevention System (IPS) is a security system that not only detects but also prevents an attack in real time. It monitors network traffic and takes action to block threats. IPS helps stop malicious activities before they can cause damage. It is an advanced form of network security protection."},
{"q":"What is hashing?","keywords":["convert","data"],"answer":"Hashing is a process used to convert data into a fixed-length value called a hash. It is commonly used for storing passwords securely. Even a small change in data produces a completely different hash value. Hashing ensures data integrity and security in many systems."},
{"q":"What is SSL?","keywords":["secure","web"],"answer":"SSL (Secure Sockets Layer) is a protocol used to secure communication on the web. It encrypts data exchanged between a client and a server. This ensures that sensitive information like passwords and credit card details remain protected. SSL is essential for maintaining trust and security in online transactions."},
{"q":"What is brute force attack?","keywords":["guess","password"],"answer":"A brute force attack is a method where attackers try to guess a password by testing many possible combinations. It relies on repeated attempts until the correct password is found. This type of attack can be prevented using strong passwords and account lock mechanisms. It is a common threat in cybersecurity."},
{"q":"What is social engineering?","keywords":["human","attack"],"answer":"Social engineering is a type of attack that targets human behavior instead of technical vulnerabilities. Attackers manipulate people into revealing sensitive information. This human-based attack can involve phishing, pretexting, or impersonation. Awareness and training are important to prevent such attacks."}
],

"advanced":[
{"q":"What is zero-day attack?","keywords":["unknown","vulnerability"],"answer":"A zero-day attack is a cyber attack that exploits an unknown vulnerability in software or systems. Since the vulnerability is not yet discovered or fixed, attackers can use it before developers respond. These attacks are highly dangerous because there is no immediate defense available. Detecting and patching such vulnerabilities quickly is critical for security."},
{"q":"What is SIEM?","keywords":["monitor","logs"],"answer":"SIEM (Security Information and Event Management) is a system used to monitor and analyze security logs in real time. It collects data from multiple sources across a network. By analyzing logs, it helps detect suspicious activities and potential threats. SIEM tools are widely used in organizations for improving security monitoring."},
{"q":"What is threat modeling?","keywords":["identify","risk"],"answer":"Threat modeling is a process used to identify potential security risk in a system. It helps developers understand possible threats and vulnerabilities. By analyzing risks early, better security measures can be implemented. This approach improves overall system security and reduces chances of attacks."},
{"q":"What is digital forensics?","keywords":["investigation","data"],"answer":"Digital forensics is the process of investigation and analysis of digital data after a cyber incident. It involves collecting and examining evidence from computers, networks, or devices. This helps in understanding how an attack happened and who was responsible. Digital forensics is important for legal and security purposes."},
{"q":"What is endpoint security?","keywords":["device","protect"],"answer":"Endpoint security focuses on protecting each device connected to a network, such as computers and smartphones. It ensures that every device is secure from threats and unauthorized access. This approach helps protect sensitive data at the device level. Endpoint security is essential in modern organizations with multiple connected devices."},
{"q":"What is cloud security?","keywords":["protect","cloud"],"answer":"Cloud security refers to the practices and technologies used to protect data and systems in the cloud. It ensures that cloud environments remain safe from cyber threats. Security measures include encryption, access control, and monitoring. Cloud security is important as more organizations store data in cloud platforms."},
{"q":"What is ransomware?","keywords":["lock","data"],"answer":"Ransomware is a type of malware that is used to lock data or systems until a ransom is paid. It prevents users from accessing their files or devices. Attackers demand payment to restore access. Protecting systems from ransomware requires strong security measures and regular backups."},
{"q":"What is DDoS attack?","keywords":["traffic","server"],"answer":"A DDoS (Distributed Denial of Service) attack occurs when multiple systems send excessive traffic to a server. This overload causes the server to slow down or crash. As a result, legitimate users cannot access the service. Preventing such attacks requires strong network security and monitoring."},
{"q":"What is identity management?","keywords":["access","control"],"answer":"Identity management is the process of managing user identities and controlling their access to systems. It ensures that only authorized users can access specific resources. This includes authentication and permission management. Proper identity management improves security and reduces unauthorized access."},
{"q":"What is security compliance?","keywords":["rules","standards"],"answer":"Security compliance refers to following established rules and standards to ensure system security. Organizations must meet legal and industry requirements for data protection. Compliance helps reduce risks and improve trust with users. It is essential for maintaining a secure and reliable system."}
]
},

"Web Developer": {

"basic":[
{"q":"What is HTML?","keywords":["structure","web"],"answer":"HTML is a markup language used to define the structure of a web page. It uses elements like headings, paragraphs, and links to organize content. This structure helps browsers display information properly on the web. HTML is the foundation of all websites and works with other technologies."},
{"q":"What is CSS?","keywords":["style","design"],"answer":"CSS is used to add style and design to web pages created using HTML. It controls colors, fonts, layouts, and spacing of elements. CSS makes web pages visually appealing and user-friendly. It plays an important role in improving the overall look and feel of websites."},
{"q":"What is JavaScript?","keywords":["interactive","web"],"answer":"JavaScript is a programming language used to make web pages interactive and dynamic. It allows developers to add features like animations, forms, and real-time updates. This interactivity improves user experience on the web. JavaScript works together with HTML and CSS."},
{"q":"What is browser?","keywords":["internet","web"],"answer":"A browser is a software application used to access and view websites on the internet. It retrieves and displays content from the web using protocols like HTTP. Popular browsers include Chrome and Firefox. Browsers help users navigate and interact with online content easily."},
{"q":"What is URL?","keywords":["address","web"],"answer":"A URL (Uniform Resource Locator) is the address used to locate resources on the web. It specifies the location of a webpage or file on the internet. Users enter a URL in a browser to access a website. It is an essential part of how the web works."},
{"q":"What is HTTP?","keywords":["protocol","web"],"answer":"HTTP (HyperText Transfer Protocol) is a protocol used for transferring data over the web. It defines how messages are sent between clients and servers. When a user accesses a website, HTTP is used to request and deliver content. It is the foundation of communication on the web."},
{"q":"What is DOM?","keywords":["document","structure"],"answer":"The DOM (Document Object Model) represents the document structure of a web page. It allows programming languages like JavaScript to access and modify content dynamically. The DOM organizes elements as a tree structure. This makes it easier to interact with and update the document."},
{"q":"What is responsive design?","keywords":["mobile","layout"],"answer":"Responsive design is an approach used to create web pages that adapt to different devices. It ensures that the layout adjusts properly for mobile, tablet, and desktop screens. This improves user experience across all devices. Responsive design uses flexible layouts and media queries."},
{"q":"What is frontend?","keywords":["ui","user"],"answer":"Frontend refers to the part of a website that is visible to the user. It includes the UI elements like buttons, forms, and layouts. Frontend development focuses on improving user interaction and experience. It is built using technologies like HTML, CSS, and JavaScript."},
{"q":"What is backend?","keywords":["server","logic"],"answer":"Backend refers to the server-side part of a web application. It handles business logic, data processing, and database interactions. The backend ensures that the application works correctly behind the scenes. It communicates with the frontend to deliver data to users."}
],

"intermediate":[
{"q":"What is API?","keywords":["communication","data"],"answer":"An API (Application Programming Interface) is used for communication between different software systems. It allows the frontend and backend to exchange data efficiently. APIs define rules and endpoints that applications follow to interact. They play a crucial role in modern web development."},
{"q":"What is JSON?","keywords":["data","format"],"answer":"JSON (JavaScript Object Notation) is a lightweight data format used for exchanging data between systems. It is easy to read and write for both humans and machines. JSON represents data in key-value pairs, making it structured and simple. It is widely used in APIs and web applications."},
{"q":"What is React?","keywords":["library","ui"],"answer":"React is a JavaScript library used to build user interfaces. It allows developers to create reusable UI components. React improves performance by updating only the necessary parts of the UI. It is widely used for building modern and dynamic web applications."},
{"q":"What is Node.js?","keywords":["javascript","server"],"answer":"Node.js is a runtime environment that allows JavaScript to run on the server. It enables developers to use javascript for backend development. Node.js is fast and efficient due to its non-blocking architecture. It is commonly used for building scalable web applications."},
{"q":"What is Express?","keywords":["framework","node"],"answer":"Express is a web framework built on Node that simplifies backend development. It provides tools for handling routes, requests, and responses. Express makes it easier to build APIs and web servers. It is widely used in combination with Node.js."},
{"q":"What is routing?","keywords":["url","navigation"],"answer":"Routing is the process of mapping a url to a specific function or page in a web application. It controls navigation between different parts of a website. Routing ensures that users see the correct content based on the url they access. It is an essential part of web development."},
{"q":"What is middleware?","keywords":["process","request"],"answer":"Middleware is a function that processes a request before it reaches the final handler. It can modify request and response objects or perform tasks like authentication. Middleware helps manage the flow of data in an application. It is widely used in backend frameworks like Express."},
{"q":"What is session?","keywords":["user","data"],"answer":"A session is used to store user data temporarily on the server during a user's interaction. It helps maintain user state across multiple requests. Sessions are commonly used for login systems and authentication. They ensure a smooth and personalized user experience."},
{"q":"What is cookie?","keywords":["browser","data"],"answer":"A cookie is a small piece of data stored in the browser. It is used to remember user preferences and session information. Cookies help websites recognize returning users. They play an important role in maintaining user sessions and tracking data."},
{"q":"What is AJAX?","keywords":["async","request"],"answer":"AJAX (Asynchronous JavaScript and XML) is a technique used to send async request to the server without reloading the page. It allows web applications to update data dynamically. AJAX improves user experience by making applications faster and more interactive. It is widely used in modern web development."}
],

"advanced":[
{"q":"What is SPA?","keywords":["single","page"],"answer":"A SPA (Single Page Application) is a type of web application that loads a single page and updates content dynamically. Instead of reloading the entire page, it changes only the required parts. This improves performance and provides a smoother user experience. SPAs are commonly built using modern JavaScript frameworks."},
{"q":"What is SSR?","keywords":["server","render"],"answer":"SSR (Server Side Rendering) is a technique where web pages are rendered on the server before being sent to the client. This improves initial load time and SEO performance. The server prepares the HTML content and sends it to the browser. SSR is useful for applications that need fast and optimized rendering."},
{"q":"What is JWT?","keywords":["token","auth"],"answer":"JWT (JSON Web Token) is a method used for secure authentication between client and server. It uses a token that contains encoded user information. This token is sent with each request to verify the user's identity. JWT is widely used in modern web applications for stateless authentication."},
{"q":"What is web security?","keywords":["protect","web"],"answer":"Web security involves techniques and practices used to protect web applications from cyber threats. It ensures that data and systems remain safe from attacks. Common measures include encryption, authentication, and input validation. Web security is essential for maintaining trust and safety online."},
{"q":"What is CDN?","keywords":["delivery","content"],"answer":"A CDN (Content Delivery Network) is a network of servers that helps in faster delivery of web content. It stores copies of content at different locations around the world. This reduces load time by serving data from the nearest server. CDN improves performance and user experience."},
{"q":"What is lazy loading?","keywords":["performance","load"],"answer":"Lazy loading is a technique used to improve performance by loading content only when it is needed. Instead of loading all resources at once, it loads them as the user scrolls or interacts. This reduces initial load time and saves bandwidth. Lazy loading is commonly used for images and videos."},
{"q":"What is webpack?","keywords":["bundle","assets"],"answer":"Webpack is a tool used to bundle JavaScript files and other assets like CSS and images. It combines multiple files into a single bundle for better performance. Webpack helps optimize code and manage dependencies. It is widely used in modern frontend development."},
{"q":"What is SEO?","keywords":["search","optimize"],"answer":"SEO (Search Engine Optimization) is the process used to optimize websites for better search engine ranking. It helps improve visibility in search results. Techniques include keyword optimization, content improvement, and performance enhancements. SEO is important for increasing website traffic."},
{"q":"What is PWA?","keywords":["app","web"],"answer":"PWA (Progressive Web App) is a type of web application that behaves like a mobile app. It works on web browsers but provides features like offline access and notifications. PWAs improve user experience and performance. They combine the best features of web and mobile apps."},
{"q":"What is web socket?","keywords":["real-time","communication"],"answer":"WebSocket is a protocol that enables real-time communication between client and server. It allows continuous data exchange without repeatedly sending requests. This makes it faster and more efficient than traditional HTTP. WebSockets are used in chat apps, live updates, and gaming."}
]
},

"Cloud Engineer": {

"basic":[
{"q":"What is cloud computing?","keywords":["internet","services"],"answer":"Cloud computing is a technology that provides computing services over the internet. These services include storage, servers, databases, and networking. Instead of using local systems, users can access resources online. Cloud computing offers flexibility, scalability, and cost efficiency."},
{"q":"What is AWS?","keywords":["amazon","cloud"],"answer":"AWS (Amazon Web Services) is a popular cloud platform provided by Amazon. It offers a wide range of cloud services like storage, computing, and databases. AWS allows businesses to build and deploy applications easily. It is widely used for scalable and reliable cloud solutions."},
{"q":"What is storage?","keywords":["data","save"],"answer":"Storage in cloud computing refers to saving data in remote servers instead of local devices. It allows users to store, manage, and access data from anywhere. Cloud storage ensures data safety and backup. It is an essential part of modern applications."},
{"q":"What is VM?","keywords":["virtual","machine"],"answer":"A VM (Virtual Machine) is a virtual version of a physical computer. It runs on a host system and behaves like a real machine. VMs are used to run multiple operating systems on a single device. They are widely used in cloud environments for flexibility."},
{"q":"What is scalability?","keywords":["increase","load"],"answer":"Scalability is the ability of a system to handle an increase in load efficiently. It allows systems to add more resources when needed. This ensures smooth performance even during high demand. Scalability is a key feature of cloud computing."},
{"q":"What is server?","keywords":["computer","network"],"answer":"A server is a powerful computer that provides services to other devices on a network. It handles requests and delivers data to clients. Servers are essential for hosting websites, applications, and databases. They play a central role in network communication."},
{"q":"What is backup?","keywords":["data","copy"],"answer":"Backup is the process of creating a copy of data to prevent data loss. It ensures that information can be restored in case of failure or attack. Backups are stored in secure locations. Regular backups are important for data safety."},
{"q":"What is region?","keywords":["location","cloud"],"answer":"A region in cloud computing refers to a specific geographical location where data centers are located. Each region contains multiple availability zones. Choosing the right region helps improve performance and reduce latency. It is important for data management and compliance."},
{"q":"What is uptime?","keywords":["availability","time"],"answer":"Uptime refers to the amount of time a system or service is available and operational. It is usually expressed as a percentage. High uptime ensures reliability and continuous access for users. It is a key factor in evaluating system performance."},
{"q":"What is SLA?","keywords":["agreement","service"],"answer":"SLA (Service Level Agreement) is a formal agreement between a provider and a user. It defines the level of service expected, including performance and uptime. SLAs ensure accountability and reliability. They are important in cloud and IT services."}
],

"intermediate":[
{"q":"What is IaaS?","keywords":["infrastructure","cloud"],"answer":"IaaS (Infrastructure as a Service) is a cloud computing model that provides virtualized infrastructure over the cloud. It includes resources like servers, storage, and networking. Users can manage and control the infrastructure without maintaining physical hardware. It offers flexibility and scalability for businesses."},
{"q":"What is PaaS?","keywords":["platform","cloud"],"answer":"PaaS (Platform as a Service) is a cloud service that provides a platform for developers to build, test, and deploy applications. It includes tools, frameworks, and environments needed for development. Developers do not need to manage underlying infrastructure. PaaS simplifies the application development process."},
{"q":"What is SaaS?","keywords":["software","service"],"answer":"SaaS (Software as a Service) is a cloud model that delivers software applications over the internet. Users can access the software through a browser without installing it. The service provider manages updates and maintenance. SaaS is widely used for applications like email and online tools."},
{"q":"What is load balancer?","keywords":["traffic","server"],"answer":"A load balancer is a system that distributes incoming traffic across multiple server resources. It ensures that no single server becomes overloaded. This improves performance, reliability, and availability of applications. Load balancing is essential for handling high user demand."},
{"q":"What is auto scaling?","keywords":["adjust","resources"],"answer":"Auto scaling is a feature in cloud computing that automatically adjusts resources based on demand. It increases resources when traffic is high and reduces them when demand is low. This helps optimize cost and performance. Auto scaling ensures efficient use of cloud resources."},
{"q":"What is container?","keywords":["app","package"],"answer":"A container is a lightweight unit that packages an app along with its dependencies. It ensures that the app runs consistently across different environments. Containers are portable and efficient. They are widely used in modern application development and deployment."},
{"q":"What is Docker?","keywords":["container","tool"],"answer":"Docker is a popular tool used to create, manage, and run container-based applications. It helps developers package applications into containers. Docker ensures consistency across development and production environments. It simplifies the deployment process."},
{"q":"What is Kubernetes?","keywords":["manage","containers"],"answer":"Kubernetes is a platform used to manage and orchestrate containers at scale. It automates deployment, scaling, and monitoring of containerized applications. Kubernetes helps ensure high availability and efficient resource usage. It is widely used in cloud-native applications."},
{"q":"What is cloud security?","keywords":["protect","data"],"answer":"Cloud security involves strategies and technologies used to protect data in cloud environments. It includes encryption, access control, and monitoring. These measures ensure that sensitive information remains safe. Cloud security is essential for maintaining trust and compliance."},
{"q":"What is monitoring?","keywords":["track","system"],"answer":"Monitoring is the process used to track system performance and health in real time. It helps detect issues, errors, or failures early. Monitoring tools provide alerts and reports for better management. It is important for maintaining reliable and efficient systems."}
],

"advanced":[
{"q":"What is multi-cloud?","keywords":["multiple","cloud"],"answer":"Multi-cloud is a strategy where organizations use multiple cloud providers instead of relying on a single one. This approach increases flexibility and reduces dependency on one vendor. It also improves reliability and performance by distributing workloads. Using multiple cloud platforms helps optimize cost and services."},
{"q":"What is hybrid cloud?","keywords":["private","public"],"answer":"Hybrid cloud is a cloud computing model that combines private and public cloud environments. It allows data and applications to be shared between them. This provides flexibility, better control, and improved security. Organizations can keep sensitive data in private cloud while using public cloud for scalability."},
{"q":"What is cloud architecture?","keywords":["design","system"],"answer":"Cloud architecture refers to the design and structure of cloud-based systems. It includes components like servers, storage, networking, and services. A well-designed system ensures performance, scalability, and reliability. Cloud architecture is important for building efficient applications."},
{"q":"What is fault tolerance?","keywords":["failure","recover"],"answer":"Fault tolerance is the ability of a system to continue working even when a failure occurs. It ensures that systems can recover quickly without affecting users. This is achieved using backup systems and redundancy. Fault tolerance is important for maintaining reliability and uptime."},
{"q":"What is disaster recovery?","keywords":["restore","data"],"answer":"Disaster recovery is the process of restoring data and systems after a failure or disaster. It ensures business continuity in case of data loss or system crash. Recovery plans include backups and failover systems. It is essential for protecting critical data."},
{"q":"What is serverless?","keywords":["no","server"],"answer":"Serverless computing is a cloud model where developers can run code without managing server infrastructure. Even though servers exist, users do not need to handle them. The cloud provider manages scaling and resources automatically. This makes development faster and more efficient."},
{"q":"What is edge computing?","keywords":["near","data"],"answer":"Edge computing is a method of processing data near the source where it is generated. This reduces latency and improves performance. Instead of sending all data to a central server, processing happens closer to devices. It is useful in applications like IoT and real-time systems."},
{"q":"What is cost optimization?","keywords":["reduce","cost"],"answer":"Cost optimization in cloud computing focuses on reducing unnecessary expenses while maintaining performance. It involves choosing the right resources and scaling efficiently. Monitoring usage helps identify areas to save cost. This ensures better utilization of cloud services."},
{"q":"What is IAM?","keywords":["access","control"],"answer":"IAM (Identity and Access Management) is a system used to manage user identities and control their access to resources. It ensures that only authorized users can access specific services. IAM includes authentication and permission management. It is essential for maintaining security in cloud systems."},
{"q":"What is logging?","keywords":["record","activity"],"answer":"Logging is the process of keeping a record of system activity and events. It helps track user actions, errors, and system performance. Logs are useful for debugging and security monitoring. Proper logging ensures better management and analysis of systems."}
]
},

"DevOps Engineer": {

"basic":[
{"q":"What is DevOps?","keywords":["development","operations"],"answer":"DevOps is a practice that combines development and operations teams to improve collaboration. It focuses on faster delivery of software and continuous improvement. By integrating development and operations, it ensures better efficiency and reliability. DevOps uses automation and monitoring to streamline processes."},
{"q":"What is CI/CD?","keywords":["automation","deploy"],"answer":"CI/CD stands for Continuous Integration and Continuous Deployment, which focuses on automation in software development. It helps automate the process of testing and deploy code changes. This ensures faster delivery and fewer errors. CI/CD pipelines improve efficiency and maintain code quality."},
{"q":"What is Git?","keywords":["version","control"],"answer":"Git is a version control system used to track changes in code over time. It allows developers to collaborate and manage code efficiently. With version control, previous versions of code can be restored easily. Git is widely used in software development projects."},
{"q":"What is build?","keywords":["compile","code"],"answer":"A build is the process of converting source code into an executable program. It involves compiling code and combining different components. Builds ensure that the code works correctly before deployment. This step is important in the software development lifecycle."},
{"q":"What is deployment?","keywords":["release","app"],"answer":"Deployment is the process of releasing an app to users after development and testing. It involves moving the application to a production environment. Proper deployment ensures that the app runs smoothly. It is a key step in delivering software to users."},
{"q":"What is automation?","keywords":["auto","process"],"answer":"Automation is the use of tools and scripts to perform tasks automatically without manual effort. It helps speed up repetitive processes and reduce errors. Automation is widely used in testing, deployment, and monitoring. It improves efficiency and productivity in development workflows."},
{"q":"What is testing?","keywords":["check","code"],"answer":"Testing is the process used to check code for errors and ensure it works as expected. It involves running different test cases on the application. Testing helps improve quality and reliability of software. It is an essential part of development."},
{"q":"What is pipeline?","keywords":["steps","process"],"answer":"A pipeline is a series of automated steps used in the software development process. It includes stages like build, test, and deployment. Pipelines help streamline the workflow and reduce manual effort. They are commonly used in CI/CD systems."},
{"q":"What is script?","keywords":["code","automation"],"answer":"A script is a small program written in code to automate tasks. It helps perform repetitive actions quickly and efficiently. Scripts are commonly used in automation processes. They save time and reduce manual work."},
{"q":"What is environment?","keywords":["setup","system"],"answer":"An environment refers to the setup of a system where applications are developed or run. It includes software, hardware, and configurations. Different environments like development, testing, and production are used. Proper setup ensures smooth functioning of applications."}
],

"intermediate":[
{"q":"What is Jenkins?","keywords":["ci","tool"],"answer":"Jenkins is a popular ci tool used for automating software development processes. It helps in building, testing, and deploying code automatically. Jenkins supports continuous integration and continuous delivery pipelines. It improves development speed and reduces manual effort."},
{"q":"What is Docker?","keywords":["container","app"],"answer":"Docker is a platform used to create and run applications inside a container. It packages an app along with its dependencies to ensure consistency. This makes it easy to run the app in different environments. Docker simplifies deployment and improves portability."},
{"q":"What is Kubernetes?","keywords":["manage","containers"],"answer":"Kubernetes is a platform used to manage containerized applications at scale. It automates deployment, scaling, and monitoring of containers. Kubernetes ensures high availability and efficient resource usage. It is widely used in modern cloud environments."},
{"q":"What is Ansible?","keywords":["automation","tool"],"answer":"Ansible is an automation tool used for configuration management and application deployment. It helps automate repetitive tasks without complex setup. Ansible uses simple scripts called playbooks. It improves efficiency and consistency in managing systems."},
{"q":"What is Terraform?","keywords":["infrastructure","code"],"answer":"Terraform is a tool used to manage infrastructure using code. It allows users to define and provision resources in a cloud environment. Infrastructure as code helps automate setup and scaling. Terraform improves consistency and reduces manual configuration errors."},
{"q":"What is monitoring?","keywords":["track","system"],"answer":"Monitoring is the process used to track system performance and health continuously. It helps detect issues and errors in real time. Monitoring tools provide alerts and insights for better system management. It ensures reliability and smooth operation of systems."},
{"q":"What is logging?","keywords":["record","data"],"answer":"Logging is the process of keeping a record of system data and events. It helps in tracking activities and debugging issues. Logs provide useful information about system behavior. Proper logging is important for security and performance analysis."},
{"q":"What is scaling?","keywords":["increase","resources"],"answer":"Scaling is the process of increasing or decreasing resources based on system demand. It helps handle more users or reduce cost during low usage. Scaling can be vertical or horizontal. It ensures efficient use of resources and better performance."},
{"q":"What is containerization?","keywords":["package","app"],"answer":"Containerization is a technique used to package an app along with its dependencies into a container. This ensures the app runs consistently across different environments. It improves portability and efficiency. Containerization is widely used in modern development practices."},
{"q":"What is orchestration?","keywords":["manage","containers"],"answer":"Orchestration is the process used to manage multiple containers in a system. It automates tasks like deployment, scaling, and networking of containers. Orchestration tools ensure smooth operation of applications. It is essential for handling large-scale containerized systems."}
],

"advanced":[
{"q":"What is blue-green deployment?","keywords":["release","strategy"],"answer":"Blue-green deployment is a release strategy used to reduce downtime and risk during application updates. It involves maintaining two environments, one active and one idle. The new version is deployed to the idle environment and tested before switching traffic. This strategy ensures a smooth and reliable release process."},
{"q":"What is canary deployment?","keywords":["gradual","release"],"answer":"Canary deployment is a method where a new version of an application is released gradually to a small group of users. This gradual release helps identify issues before full deployment. If everything works well, the update is rolled out to all users. It reduces risk and improves reliability."},
{"q":"What is microservices?","keywords":["services","small"],"answer":"Microservices is an architecture where an application is divided into small independent services. Each service handles a specific function and can be developed separately. These services communicate with each other using APIs. This approach improves scalability, flexibility, and maintainability."},
{"q":"What is observability?","keywords":["monitor","system"],"answer":"Observability is the ability to monitor and understand the internal state of a system. It uses metrics, logs, and traces to analyze performance. Observability helps detect issues and improve system reliability. It is important for maintaining complex applications."},
{"q":"What is incident management?","keywords":["handle","issues"],"answer":"Incident management is the process used to handle unexpected issues or failures in a system. It focuses on quickly identifying, analyzing, and resolving problems. The goal is to restore normal operation as soon as possible. Effective incident management improves system reliability."},
{"q":"What is SRE?","keywords":["reliability","engineer"],"answer":"SRE (Site Reliability Engineering) is a practice that focuses on maintaining system reliability and performance. It combines software engineering with operations tasks. An SRE engineer works to automate processes and reduce system failures. SRE ensures high availability of applications."},
{"q":"What is versioning?","keywords":["update","code"],"answer":"Versioning is the practice of tracking different versions of code during development. It helps manage updates and changes effectively. Each version represents a specific state of the application. Versioning makes it easier to maintain and improve software over time."},
{"q":"What is rollback?","keywords":["revert","deploy"],"answer":"Rollback is the process of reverting a system to a previous stable state after a failed deploy. It helps quickly recover from errors or issues. Rollback ensures that users are not affected by faulty updates. It is an important safety mechanism in deployment."},
{"q":"What is security in DevOps?","keywords":["protect","pipeline"],"answer":"Security in DevOps focuses on protecting the entire development and deployment pipeline. It includes practices like code scanning, secure configurations, and access control. Integrating security into the pipeline ensures early detection of vulnerabilities. This approach improves overall system safety."},
{"q":"What is automation testing?","keywords":["auto","test"],"answer":"Automation testing is the process of using tools to automatically test applications. It reduces manual effort and increases testing speed. Automated tests ensure code quality and reliability. It is widely used in CI/CD pipelines for continuous testing."}
]
}
}

EXPECTED_ANSWERS = {}

for role in INTERVIEW_QUESTIONS:
    for level in INTERVIEW_QUESTIONS[role]:
        for item in INTERVIEW_QUESTIONS[role][level]:
            if isinstance(item, dict):
                EXPECTED_ANSWERS[item["q"]] = item["answer"]

def evaluate_answer(user_answer, expected_keywords):

    user_answer = user_answer.lower()

    matched = []
    missing = []

    for keyword in expected_keywords:
        if keyword.lower() in user_answer:
            matched.append(keyword)
        else:
            missing.append(keyword)

    score = int((len(matched) / len(expected_keywords)) * 100) if expected_keywords else 0

    # NEW FEEDBACK SYSTEM
    if score >= 80:
        feedback = "Excellent answer. You covered most key concepts clearly."
    elif score >= 50:
        feedback = "Good attempt, but you missed some important points."
    else:
        feedback = "Weak answer. You need to focus on core concepts."

    return {
        "score": score,
        "matched": matched,
        "missing": missing,
        "feedback": feedback
    }

# Routes

@app.route("/")
def home():
    if "user_id" in session:
        return redirect(url_for("dashboard"))
    return render_template("auth.html")

@app.route("/signup", methods=["POST"])
def signup():
    username = request.form["username"].strip()
    password = request.form["password"].strip()

    existing = User.query.filter_by(username=username).first()
    if existing:
        flash("Username already exists.")
        return redirect(url_for("home"))

    hashed_password = generate_password_hash(password)
    user = User(username=username, password=hashed_password)
    db.session.add(user)
    db.session.commit()

    session["user_id"] = user.id
    session["username"] = user.username

    save_activity(user.id, "auth", "Signed up", f"User {username} created an account")
    return redirect(url_for("dashboard"))

@app.route("/login", methods=["POST"])
def login():
    username = request.form["username"].strip()
    password = request.form["password"].strip()

    user = User.query.filter_by(username=username).first()
    if user and check_password_hash(user.password, password):
        session["user_id"] = user.id
        session["username"] = user.username
        save_activity(user.id, "auth", "Logged in", f"User {username} logged in")
        return redirect(url_for("dashboard"))

    flash("Invalid username or password.")
    return redirect(url_for("home"))

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("home"))

@app.route("/dashboard")
def dashboard():
    if not login_required():
        return redirect(url_for("home"))
    return render_template("index.html", username=session.get("username"))

@app.route("/my_dashboard")
def my_dashboard():
    if not login_required():
        return redirect(url_for("home"))

    user_id = session["user_id"]
    activities = Activity.query.filter_by(user_id=user_id).order_by(Activity.timestamp.desc()).all()
    interviews = InterviewSession.query.filter_by(user_id=user_id).order_by(InterviewSession.created_at.asc()).all()

    progress_labels = [i.created_at.strftime("%d %b") for i in interviews]
    progress_scores = [i.score for i in interviews]

    total_interviews = len(interviews)
    avg_score = int(sum(progress_scores) / total_interviews) if total_interviews > 0 else 0

    return render_template(
        "user_dashboard.html",
        username=session.get("username"),
        activities=activities,
        interviews=interviews,
        progress_labels=json.dumps(progress_labels),
        progress_scores=json.dumps(progress_scores),
        total_interviews=total_interviews,
        avg_score=avg_score
    )

@app.route("/profile", methods=["GET","POST"])
def profile():
    if not login_required():
        return redirect(url_for("home"))

    user = User.query.get(session["user_id"])
    return render_template("profile.html", user=user)

@app.route("/analyze", methods=["GET", "POST"])
def analyze():
    if not login_required():
        return redirect(url_for("home"))

    results = None

    

    if request.method == "POST":
        file = request.files.get("resume")
        role = request.form.get("role")

        if file and file.filename.endswith(".pdf"):

            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            file.save(filepath)

            #  EXTRACT TEXT
            text = extract_text_from_pdf(filepath)

            #  SAFETY FIX (IMPORTANT)
            if not text or len(text.strip()) == 0:
                text = ""

            print("TEXT SAMPLE:", text[:500])

            #  ANALYZE
            results = analyze_resume_text(text, role)
            print("FULL RESULTS:", results)

            session["last_analysis"] = results

            print("SUGGESTIONS:", results.get("suggestions"))

            #  SAVE ACTIVITY
            save_activity(
                session["user_id"],
                "analyze",
                f"Resume analyzed for {role}",
                f"ATS Score: {results.get('score',0)}%"
            )

    return render_template("analyzer.html", results=results)

def analyze_resume_text(text, role):

    import re

    text_original = text   # keep original for formatting check

    # CLEAN TEXT
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text)

    # KEYWORDS
    
    ROLE_KEYWORDS = {
        "Software Developer": ["python","java","c++","javascript","react","node","git","docker","aws","api","sql"],
        "Data Analyst": ["python","sql","tableau","power bi","excel","statistics","pandas","numpy"],
        "AI Engineer": ["machine learning","deep learning","tensorflow","pytorch","nlp"],
        "Web Developer": ["html","css","javascript","react","node","mongodb"],
        "Cloud Engineer": ["aws","azure","gcp","docker","kubernetes"],
        "Cyber Security": ["network security","penetration testing","ethical hacking"],
        "DevOps Engineer": ["docker","kubernetes","jenkins","ci cd","terraform"]
    }

    keywords = ROLE_KEYWORDS.get(role, [])

    found_keywords = [k for k in keywords if k.lower() in text]
    missing_keywords = [k for k in keywords if k.lower() not in text]

    keyword_score = int((len(found_keywords)/len(keywords))*40) if keywords else 0

    # SECTIONS (IMPROVED)
    
    section_patterns = {
        "Summary": ["summary", "profile"],
        "Skills": ["skills", "technical skills"],
        "Experience": ["experience", "work experience"],
        "Education": ["education"],
        "Projects": ["projects"]
    }

    found_sections = []

    for section, variants in section_patterns.items():
        if any(v in text for v in variants):
            found_sections.append(section)

    missing_sections = [s for s in section_patterns if s not in found_sections]

    section_score = int((len(found_sections)/len(section_patterns))*20)

    #  FORMATTING
    
    formatting_score = 0

    if "-" in text_original or "•" in text_original:
        formatting_score += 5

    if len(text_original) > 300:
        formatting_score += 5

    if "\n" in text_original:
        formatting_score += 5

    #  CONTACT
    
    contact_score = 0

    if re.search(r"\S+@\S+\.\S+", text_original):
        contact_score += 5

    if re.search(r"\d{10}", text_original):
        contact_score += 5
 
    #  ACTION WORDS
    
    action_words = [
        "develop", "build", "design", "implement",
        "create", "lead", "improve", "optimize"
    ]

    found_actions = []

    for word in action_words:
        if re.search(rf"\b{word}\w*\b", text):
            found_actions.append(word)

    action_score = int((len(found_actions)/len(action_words))*15)

    # FINAL SCORE
    total_score = keyword_score + section_score + formatting_score + contact_score + action_score

    # SUGGESTIONS (FIXED)

    suggestions = []

    if missing_keywords:
        suggestions.append("Add more relevant keywords for the selected role.")

    if missing_sections:
        suggestions.append("Add missing sections: " + ", ".join(missing_sections))

    if not found_actions:
        suggestions.append("Use strong action verbs like Developed, Built, Designed.")

    if formatting_score < 10:
        suggestions.append("Improve formatting (use bullet points, spacing).")

    if contact_score < 10:
        suggestions.append("Include proper email and phone number.")

#  FORCE AT LEAST 1 SUGGESTION
    if not suggestions:
        suggestions.append("Great resume! Just fine-tune formatting and add more impact.")

    print("DEBUG SUGGESTIONS:", suggestions)

    return {
        "role": role,
        "score": total_score,
        "found": found_keywords,
        "missing": missing_keywords,
        "sections_found": found_sections,
        "missing_sections": missing_sections,
        "action_words_used": found_actions,
        "suggestions": suggestions
    }

@app.route("/builder", methods=["GET"])
def builder():
    if not login_required():
        return redirect(url_for("home"))

    #  Step 1: Get data from session
    session_data = session.get("resume_data", {})

    #  Step 2: Also allow data from GET (when coming back from preview)
    query_data = request.args.to_dict()

    #  Step 3: Merge both (priority to query_data)
    data = {**session_data, **query_data}

    return render_template("builder.html", data=data)

@app.route("/preview", methods=["POST"])
def preview():

    
    data = request.form.to_dict()
    
     #  FIX LINKS
   
    linkedin = data.get("linkedin", "")
    github = data.get("github", "")

    if linkedin and not linkedin.startswith("http"):
        linkedin = "https://" + linkedin

    if github and not github.startswith("http"):
        github = "https://" + github

    data["linkedin"] = linkedin
    data["github"] = github

    
    # FORMAT DATA
    
    skills = [s.strip() for s in data.get("skills", "").split(",") if s.strip()]
    experience = [e.strip() for e in data.get("experience", "").split(",") if e.strip()]

    projects = []
    for i in range(1, 4):
        title = data.get(f"p{i}_title", "")
        date = data.get(f"p{i}_date", "")
        info = [x.strip() for x in data.get(f"p{i}_info", "").split(",") if x.strip()]

        if title:
            projects.append({
                "title": title,
                "date": date,
                "info": info
            })

    
    # HANDLE PHOTO (SAVE PROPERLY)
    

    photo_file = request.files.get("photo")
    if photo_file and photo_file.filename != "":
        from werkzeug.utils import secure_filename

        filename = secure_filename(photo_file.filename)
        filepath = os.path.join("static", "uploads", filename)

    #  MAKE SURE FOLDER EXISTS
        os.makedirs(os.path.dirname(filepath), exist_ok=True)

        photo_file.save(filepath)

    # SAVE PATH IN SESSION
        session["resume_photo"] = filepath
        session.modified = True

    #  SAVE EVERYTHING IN SESSION (MAIN FIX)
    
    session["resume_data"] = data
    session.modified = True

    # RENDER PREVIEW
    
    return render_template(
        "preview_resume.html",
        name=data.get("name"),
        role=data.get("role"),
        email=data.get("email"),
        phone=data.get("phone"),
        linkedin=linkedin,
        github=github,
        summary=data.get("summary"),
        skills=skills,
        experience=experience,
        projects=projects,
        education=data.get("education"),
        certs=data.get("certs"),
        photo=session.get("resume_photo"),
        form_data=data
    )

@app.route("/generate_resume", methods=["POST"])
def generate_resume():

    data = request.form.to_dict()

    if not data:
        data = session.get("resume_data", {})

    # PROCESS DATA
    skills = data.get("skills", "").split(",")
    experience = data.get("experience", "").split(",")

    projects = []
    for i in range(1, 4):
        title = data.get(f"p{i}_title", "")
        date = data.get(f"p{i}_date", "")
        info = data.get(f"p{i}_info", "").split(",")

        if title:
            projects.append({
                "title": title,
                "date": date,
                "info": info
            })

    # FIX: GET PHOTO PATH FROM SESSION
    import os
    photo_path = session.get("resume_photo")

    if photo_path:
        photo_path = os.path.abspath(photo_path)  #  IMPORTANT for PDF

    #  RENDER HTML
    rendered = render_template(
        "resume_template.html",
        name=data.get("name"),
        role=data.get("role"),
        email=data.get("email"),
        phone=data.get("phone"),
        linkedin=data.get("linkedin"),
        github=data.get("github"),
        summary=data.get("summary"),
        skills=skills,
        experience=experience,
        projects=projects,
        education=data.get("education"),
        certs=data.get("certs"),
        photo=photo_path, # ✅ PASS FILE PATH
    )

    # CONVERT TO PDF
    import pdfkit

    config = pdfkit.configuration(
        wkhtmltopdf=r"C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe"
    )

    options = {
        'enable-local-file-access': None,
    }

    pdf = pdfkit.from_string(rendered, False, configuration=config, options=options)

    return send_file(
        BytesIO(pdf),
        as_attachment=True,
        download_name="Professional_Resume.pdf",
        mimetype="application/pdf"
    )

def add_hyperlink(paragraph, url, text):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Style (blue + underline like LinkedIn)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

@app.route("/download_docx", methods=["POST"])
def download_docx():

    data = request.form.to_dict()

    if not data:
        data = session.get("resume_data", {})

    doc = Document()

    # PAGE MARGINS
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    #  HEADER (PHOTO + NAME)
    
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # LEFT → PHOTO
    photo_path = session.get("resume_photo")
    if photo_path:
        try:
            table.cell(0,0).paragraphs[0].add_run().add_picture(photo_path, width=Inches(1.2))
        except:
            pass

    # RIGHT → NAME + ROLE
    cell = table.cell(0,1)
    p = cell.paragraphs[0]

    run = p.add_run(data.get("name",""))
    run.bold = True
    run.font.size = Pt(18)

    p = cell.add_paragraph(data.get("role",""))
    p.runs[0].font.size = Pt(11)

    # CONTACT LINKS
    contact = cell.add_paragraph()

# Email
    email = data.get("email", "")
    if email:
        add_hyperlink(contact, f"mailto:{email}", email)
        contact.add_run(" | ")
        
    linkedin = data.get("linkedin", "")
    if linkedin:
        if not linkedin.startswith("http"):
            linkedin = "https://" + linkedin
            add_hyperlink(contact, linkedin, "LinkedIn")
            contact.add_run(" | ")

# GitHub
    github = data.get("github", "")
    if github:
        if not github.startswith("http"):
            github = "https://" + github
        add_hyperlink(contact, github, "GitHub")

    # FUNCTION FOR HEADINGS
    
    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        return p

    #  SUMMARY
    
    add_heading("PROFESSIONAL SUMMARY")
    doc.add_paragraph(data.get("summary",""))

    
    # SKILLS
    
    add_heading("TECHNICAL SKILLS")

    skills = data.get("skills","").split(",")
    for s in skills:
        if s.strip():
            doc.add_paragraph(s.strip(), style='List Bullet')


    # EXPERIENCE

    add_heading("PROFESSIONAL EXPERIENCE")

    experience = data.get("experience","").split(",")
    for e in experience:
        if e.strip():
            doc.add_paragraph(e.strip(), style='List Bullet')

    #  PROJECTS

    add_heading("PROJECTS")

    for i in range(1,4):
        title = data.get(f"p{i}_title","")
        date = data.get(f"p{i}_date","")
        info = data.get(f"p{i}_info","").split(",")

        if title:
            p = doc.add_paragraph()
            run = p.add_run(f"{title} ({date})")
            run.bold = True

            for point in info:
                if point.strip():
                    doc.add_paragraph(point.strip(), style='List Bullet')

    # EDUCATION
    
    add_heading("EDUCATION")
    doc.add_paragraph(data.get("education",""))


    #  CERTIFICATIONS

    add_heading("CERTIFICATIONS")
    doc.add_paragraph(data.get("certs",""))

    #  SAVE FILE
    
    file_path = "resume.docx"
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)

@app.route("/chatbot", methods=["GET", "POST"])
def chatbot():
    if not login_required():
        return redirect(url_for("home"))

    response = None

    if request.method == "POST":
        message = request.form.get("message")

        if message:
            response = chatbot_reply(message)

            save_activity(
                session["user_id"],
                "chatbot",
                "Asked chatbot",
                f"Question: {message} | Response: {response}"
            )

    return render_template("chatbot.html", response=response)

@app.route("/analyzer_chat", methods=["POST"])
def analyzer_chat():

    data = request.get_json()
    message = data.get("message", "").lower()

    #  GET LAST ANALYSIS
    results = session.get("last_analysis", {})

    if not results:
        return jsonify({"reply": "Please analyze your resume first."})

    #  SMART RESPONSES
    if "score" in message:
        reply = f"Your ATS score is {results.get('score', 0)}%. Try improving keywords and sections."

    elif "keyword" in message:
        found = ", ".join(results.get("found", []))
        missing = ", ".join(results.get("missing", []))
        reply = f"Found: {found}. Missing: {missing}"

    elif "improve" in message or "suggest" in message:
        suggestions = results.get("suggestions", [])
        reply = "Suggestions:\n- " + "\n- ".join(suggestions) if suggestions else "Your resume looks good!"

    elif "section" in message:
        reply = f"Missing sections: {', '.join(results.get('missing_sections', []))}"

    else:
        reply = "Ask about score, keywords, sections, or improvements."

    return jsonify({"reply": reply})

@app.route("/chatbot_api", methods=["POST"])
def chatbot_api():

    user_msg = request.json.get("message", "").lower()

    if "ats score" in user_msg:
        reply = "ATS score depends on keywords, formatting, and structure."

    elif "improve resume" in user_msg:
        reply = "Use strong action words, add projects, and include job-specific keywords."

    elif "skills" in user_msg:
        reply = "Add skills like Python, SQL, Git, and tools relevant to your target role."

    elif "projects" in user_msg:
        reply = "Include real-world projects with impact and technologies used."

    elif "experience" in user_msg:
        reply = "Write experience using bullet points with action verbs."

    else:
        reply = generate_ai_response(user_msg)

    return {"reply": reply}

@app.route("/interview")
def interview():
    if not login_required():
        return redirect(url_for("home"))
    return render_template("interview.html")

@app.route("/start_interview", methods=["POST"])
def start_interview():

    if not login_required():
        return {"error": "Unauthorized"}, 401

    role = request.json.get("role", "Software Developer")
    level = request.json.get("level", "basic")

    role_data = INTERVIEW_QUESTIONS.get(role, INTERVIEW_QUESTIONS["Software Developer"])
    questions = role_data.get(level, role_data["basic"])

    random.shuffle(questions)

    session["interview_data"] = {
        "role": role,
        "level": level,
        "questions": questions,
        "index": 0,
        "correct_answers": 0,
        "scores": [],
        "details": []
    }

    return {
        "message": "Interview started",
        "question": questions[0]["q"],
        "question_number": 1,
        "total_questions": len(questions)
    }

@app.route("/submit_interview_answer", methods=["POST"])
def submit_interview_answer():

    data = request.get_json()
    answer = data.get("answer", "").lower()

    interview_data = session.get("interview_data")

    if not interview_data:
        return jsonify({"error": "Interview session not found"})

    questions = interview_data["questions"]
    index = interview_data["index"]

    #  Current question
    question_data = questions[index]
    question_text = question_data["q"]
    expected_keywords = question_data["keywords"]

    #  FIX 1: CALL evaluation properly
    evaluation = evaluate_answer(answer, expected_keywords)

    #  FIX 2: Store details ONLY ONCE
    interview_data["details"].append({
        "question": question_text,
        "user_answer": answer,
        "expected_answer": EXPECTED_ANSWERS.get(
            question_text,
            "Answer explanation not available."
        ),
        "score": evaluation["score"],
        "matched": evaluation["matched"],
        "missing": evaluation["missing"],
        "feedback": evaluation["feedback"]
    })

    # FIX 3: scoring logic
    if evaluation["score"] >= 50:
        interview_data["correct_answers"] += 1

    # Move to next question
    interview_data["index"] += 1
    session["interview_data"] = interview_data
    session.modified = True

    # INTERVIEW COMPLETED
    
    if interview_data["index"] >= len(questions):

        total = len(questions)
        correct = interview_data["correct_answers"]
        final_score = int((correct / total) * 100)

        # Expected answers list
        expected_answers = []
        for q in questions:
            expected_answers.append({
                "question": q["q"],
                "answer": EXPECTED_ANSWERS.get(
                    q["q"],
                    "Answer explanation not available."
                )
            })

        # Save to DB
        interview_session = InterviewSession(
            user_id=session["user_id"],
            role=interview_data["role"],
            level=interview_data["level"],
            score=final_score,
            total_questions=total,
            correct_answers=correct,
            transcript=json.dumps(interview_data["details"])
        )

        db.session.add(interview_session)
        db.session.commit()

        return jsonify({
            "finished": True,
            "final_score": final_score,
            "correct_answers": correct,
            "total_questions": total,
            "expected_answers": expected_answers,
            "details": interview_data["details"],
            "last_score": evaluation["score"]
        })

    # NEXT QUESTION

    next_q = questions[interview_data["index"]]

    return jsonify({
        "finished": False,
        "next_question": next_q["q"],
        "question_number": interview_data["index"] + 1,
        "total_questions": len(questions),
        "last_score": evaluation["score"]
    })

if __name__ == "__main__":
    with app.app_context():
        db.create_all()
    app.run(host="0.0.0.0", port=5000, debug=True)